"""
PDF presentation layer for SecureChain.
Lapisan tampilan/PDF — bukan inti keamanan. Inti blockchain & hashing ada di C++.

Tugas modul ini:
  - stamp_and_secure(): beri cap visual berisi kode unik, sisipkan metadata +
    tanda tangan digital (HMAC-SHA256), lalu kunci PDF dengan password = kode unik
    (enkripsi AES-256 native PDF). Hasilnya: file hanya bisa dibuka dengan kode unik.
  - can_open(): cek apakah sebuah PDF bisa dibuka dengan kode tertentu.
  - make_sample_pdf(): buat PDF contoh untuk testing.
"""
import hmac
import hashlib
import os
import tempfile
import qrcode
from io import BytesIO
from PIL import Image

from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.colors import Color
from reportlab.lib.utils import ImageReader

# Kunci tanda tangan kampus (UBAH untuk produksi). Ini menandatangani metadata —
# membuktikan dokumen diterbitkan kampus. ponytail: HMAC = tanda tangan simetris,
# cukup untuk "bukti penerbit". Upgrade ke RSA jika butuh verifikasi pihak ketiga.
CAMPUS_SIGNING_KEY = b"SecureChain2024!Key@Campus"


def sign(student_name: str, student_id: str, code: str) -> str:
    """Tanda tangan digital HMAC-SHA256 atas data mahasiswa + kode unik."""
    msg = f"{student_name}|{student_id}|{code}".encode("utf-8")
    return hmac.new(CAMPUS_SIGNING_KEY, msg, hashlib.sha256).hexdigest()


def _make_stamp(width: float, height: float, code: str, name: str) -> PdfReader:
    """Overlay cap 1 halaman seukuran halaman target."""
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=(width, height))

    qr = qrcode.QRCode(border=1)
    qr.add_data(code)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white")
    qr_reader = ImageReader(qr_img.convert("RGB"))

    c.saveState()
    c.translate(width - 270, 40)
    c.setStrokeColor(Color(0.85, 0.05, 0.05, alpha=0.95))
    c.setFillColor(Color(0.85, 0.05, 0.05, alpha=0.12))
    c.setLineWidth(2.5)
    c.roundRect(0, 0, 250, 70, 8, stroke=1, fill=1)

    c.setFillColor(Color(0.85, 0.05, 0.05, alpha=0.98))
    c.setFont("Helvetica-Bold", 11)
    c.drawString(12, 52, "SECURECHAIN VERIFIED")
    c.setFont("Helvetica", 8)
    c.drawString(12, 38, "Kode Unik:")
    c.setFont("Helvetica-Bold", 9)
    c.drawString(12, 26, code[:24])
    c.setFont("Helvetica-Oblique", 7)
    c.drawString(12, 12, name[:28])

    c.drawImage(qr_reader, 185, 10, width=50, height=50)
    c.restoreState()

    # Watermark diagonal samar.
    c.saveState()
    c.translate(width / 2, height / 2)
    c.rotate(45)
    c.setFont("Helvetica-Bold", 46)
    c.setFillColor(Color(0.85, 0.05, 0.05, alpha=0.07))
    c.drawCentredString(0, 0, "SECURECHAIN")
    c.restoreState()

    c.save()
    buf.seek(0)
    return PdfReader(buf)


def ensure_pdf(input_path: str) -> tuple[str, bool]:
    ext = os.path.splitext(input_path)[1].lower()
    if ext == ".pdf":
        return input_path, False

    temp_dir = tempfile.gettempdir()
    temp_pdf = os.path.join(temp_dir, f"SCDV_conv_{os.path.basename(input_path)}.pdf")

    if ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp"):
        img = Image.open(input_path)
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(temp_pdf, pagesize=A4)
        width_a4, height_a4 = A4
        img_w, img_h = img.size
        ratio = min(width_a4 / img_w, height_a4 / img_h)
        new_w = img_w * ratio
        new_h = img_h * ratio
        x_offset = (width_a4 - new_w) / 2
        y_offset = (height_a4 - new_h) / 2
        c.drawImage(input_path, x_offset, y_offset, width=new_w, height=new_h)
        c.showPage()
        c.save()
        return temp_pdf, True

    elif ext == ".docx":
        import docx
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = docx.Document(input_path)
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']

        story = []
        for para in doc.paragraphs:
            if para.text.strip():
                story.append(Paragraph(para.text, normal_style))
                story.append(Spacer(1, 6))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            story.append(Paragraph(para.text, normal_style))
                            story.append(Spacer(1, 4))
                story.append(Spacer(1, 6))

        if not story:
            story.append(Paragraph("Empty Document", normal_style))

        pdf_doc = SimpleDocTemplate(temp_pdf, pagesize=A4)
        pdf_doc.build(story)
        return temp_pdf, True

    elif ext in (".txt", ".doc"):
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']

        story = []
        try:
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    line_text = line.strip()
                    if line_text:
                        story.append(Paragraph(line_text, normal_style))
                    else:
                        story.append(Spacer(1, 6))
        except Exception:
            story.append(Paragraph(f"Content of {os.path.basename(input_path)}", normal_style))

        if not story:
            story.append(Paragraph("Empty Document", normal_style))

        pdf_doc = SimpleDocTemplate(temp_pdf, pagesize=A4)
        pdf_doc.build(story)
        return temp_pdf, True

    else:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet

        styles = getSampleStyleSheet()
        normal_style = styles['Normal']

        story = [Paragraph(f"SecureChain Document: {os.path.basename(input_path)}", normal_style)]
        pdf_doc = SimpleDocTemplate(temp_pdf, pagesize=A4)
        pdf_doc.build(story)
        return temp_pdf, True


def stamp_and_secure(input_pdf: str, output_pdf: str, code: str,
                     student_name: str, student_id: str) -> str:
    """
    Beri cap + metadata + tanda tangan, lalu kunci dengan password = kode unik.
    Mengembalikan tanda tangan digital (hex). Hasil ditulis ke output_pdf.
    """
    pdf_path, is_temp = ensure_pdf(input_pdf)
    try:
        reader = PdfReader(pdf_path)
        writer = PdfWriter()

        for i, page in enumerate(reader.pages):
            if i == 0:  # cap hanya di halaman pertama (ijazah umumnya 1 halaman)
                w = float(page.mediabox.width)
                h = float(page.mediabox.height)
                page.merge_page(_make_stamp(w, h, code, student_name).pages[0])
            writer.add_page(page)

        signature = sign(student_name, student_id, code)
        writer.add_metadata({
            "/Title": "SecureChain Verified Diploma",
            "/Author": "SecureChain Campus Authority",
            "/Producer": "SecureChain Diploma Verifier (SCDV)",
            "/SCDV_Student": student_name,
            "/SCDV_StudentID": student_id,
            "/SCDV_Signature": signature,
        })

        # Enkripsi AES-256 native PDF. user_password = kode unik => file HANYA bisa
        # dibuka dengan kode unik di reader manapun (Adobe/Chrome/dll).
        writer.encrypt(user_password=code, algorithm="AES-256")

        with open(output_pdf, "wb") as f:
            writer.write(f)
    finally:
        if is_temp and os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
            except Exception:
                pass

    return signature


def can_open(pdf_path: str, code: str) -> bool:
    """True jika PDF bisa dibuka dengan kode ini."""
    try:
        reader = PdfReader(pdf_path)
        if not reader.is_encrypted:
            return True
        return reader.decrypt(code) != 0  # 0 = gagal
    except Exception:
        return False


def make_sample_pdf(path: str, name: str, student_id: str, program: str = "Teknik Informatika"):
    """Buat PDF ijazah contoh (polos, tanpa cap) untuk testing."""
    c = canvas.Canvas(path, pagesize=(595, 842))  # A4 portrait (pt)
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(297, 760, "UNIVERSITAS GADJAH MADA")
    c.setFont("Helvetica", 14)
    c.drawCentredString(297, 730, "IJAZAH SARJANA")
    c.setFont("Helvetica", 12)
    c.drawString(80, 650, f"Dengan ini menyatakan bahwa:")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(80, 615, name)
    c.setFont("Helvetica", 12)
    c.drawString(80, 585, f"NIM        : {student_id}")
    c.drawString(80, 565, f"Program  : {program}")
    c.drawString(80, 545, "Telah dinyatakan LULUS dan berhak menyandang gelar Sarjana.")
    c.setFont("Helvetica-Oblique", 10)
    c.drawString(80, 480, "Yogyakarta, 29 Juni 2026")
    c.drawString(80, 440, "Rektor,")
    c.drawString(80, 400, "Prof. Dr. Sutrisno")
    c.save()
    return path


if __name__ == "__main__":
    # Smoke test: buat -> amankan -> cek buka dengan kode benar/salah.
    import os, tempfile
    tmp = tempfile.gettempdir()
    src = os.path.join(tmp, "scdv_sample_src.pdf")
    out = os.path.join(tmp, "scdv_sample_secured.pdf")
    make_sample_pdf(src, "Hudzaifah Rahman", "010203")
    sig = stamp_and_secure(src, out, "UGM-010203-HUDZAIFAH", "Hudzaifah Rahman", "010203")
    print("signature =", sig[:24], "...")
    print("buka kode benar :", can_open(out, "UGM-010203-HUDZAIFAH"))
    print("buka kode salah :", can_open(out, "SALAH"))
    assert can_open(out, "UGM-010203-HUDZAIFAH") is True
    assert can_open(out, "SALAH") is False
    print("PDF PIPELINE OK")
